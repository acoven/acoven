#!/usr/bin/env python3
"""
Convert HTML to DOCX with ATS-friendly formatting and proper tabs.
"""

import sys
import re
import html
import subprocess
import shutil
from pathlib import Path

scriptname = "html_to_docx_ats.py"
version = "Copyright © 2025...2026-04-19.001:a@cov.in. All Rights Reserved."
SCRIPT_DIR = Path(__file__).resolve().parent

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Mm
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from bs4 import BeautifulSoup, Comment
except ImportError as e:
    print(f"Missing library: {e}")
    print("Attempting automatic install of required packages...")

    required_packages = ["python-docx", "beautifulsoup4"]
    installed = False

    # Try Homebrew first on macOS environments where pip may be externally managed.
    brew_path = shutil.which("brew")
    if brew_path:
        print("Trying Homebrew first...")
        brew_cmd = [brew_path, "install", *required_packages]
        brew_result = subprocess.run(brew_cmd, check=False)
        if brew_result.returncode == 0:
            installed = True
        else:
            print("Homebrew install failed. Falling back to pip...")

    if not installed:
        in_virtualenv = (
            hasattr(sys, "real_prefix") or
            (hasattr(sys, "base_prefix") and sys.base_prefix != sys.prefix)
        )
        pip_cmd = [sys.executable, "-m", "pip", "install"]
        if not in_virtualenv:
            pip_cmd.append("--user")
        pip_cmd.extend(required_packages)
        pip_result = subprocess.run(pip_cmd, check=False)

        # PEP 668 fallback for externally managed Python installs.
        if pip_result.returncode != 0 and not in_virtualenv:
            print("Retrying pip with --break-system-packages...")
            pip_cmd = [
                sys.executable,
                "-m",
                "pip",
                "install",
                "--user",
                "--break-system-packages",
                *required_packages,
            ]
            pip_result = subprocess.run(pip_cmd, check=False)

        if pip_result.returncode != 0:
            print("Auto-install failed.")
            print("Install manually with:")
            print("  brew install python-docx beautifulsoup4")
            print("or:")
            print("  python3 -m pip install --user python-docx beautifulsoup4")
            sys.exit(1)

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Mm
        from docx.enum.text import WD_TAB_ALIGNMENT, WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
        from docx.enum.style import WD_STYLE_TYPE
        from docx.oxml import OxmlElement
        from bs4 import BeautifulSoup, Comment
        print("Dependencies installed successfully.")
    except ImportError as retry_error:
        print(f"Import still failing after install: {retry_error}")
        sys.exit(1)

# XML namespace URLs - constants to avoid repetition
URL = {
    'schema': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'ns': 'http://www.w3.org/XML/1998/namespace'
}
URL['link'] = URL['schema'] + '/hyperlink'

# Bullet indentation constants  
BULLET_LEFT_INDENT = Inches(0.20)      # Where text wraps
BULLET_FIRST_LINE_INDENT = Inches(-0.15)  # Bullet hangs back
# Note: We want text to align at BULLET_LEFT_INDENT, so tab should be at that absolute position

# Base font size for calculations (CSS percentages and relative units)
BASE_FONT_SIZE = 10  # Base font size in points - all relative sizes calculated from this

# Page dimensions for pagination calculations
PAGE_WIDTH = 8.5  # inches (Letter size)
PAGE_HEIGHT = 11.0  # inches (Letter size)
PAGE_LEFT_MARGIN = 0.6  # inches
PAGE_RIGHT_MARGIN = 0.6  # inches
PAGE_TOP_MARGIN = 0.6  # inches
PAGE_BOTTOM_MARGIN = 0.6  # inches
USABLE_PAGE_WIDTH = PAGE_WIDTH - PAGE_LEFT_MARGIN - PAGE_RIGHT_MARGIN  # 7.3 inches
USABLE_PAGE_HEIGHT = PAGE_HEIGHT - PAGE_TOP_MARGIN - PAGE_BOTTOM_MARGIN  # 9.8 inches = 705.6 points
# Tab stops in Word are measured from the left edge of the page
# For right-aligned date tabs, compute from actual section margins at runtime.

# Centralized class pattern configuration
# These patterns are used to identify elements by function, not by specific class names
DATE_CLASS_PATTERNS = ['date']  # Class names containing these strings are treated as date elements
SEPARATOR_CLASS_PATTERNS = ['vertical-bar']  # Class names that should be replaced with separator text
SEPARATOR_REPLACEMENT = ' | '  # Text to replace separator elements with


def is_date_element(element):
    """Check if an element is likely a date element based on class patterns only"""
    if not element:
        return False
    
    classes = element.get('class', [])
    
    # Check if any class matches date patterns
    if any(any(pattern in str(c).lower() for pattern in DATE_CLASS_PATTERNS) for c in classes):
        return True
    
    return False


def get_right_tab_stop_position(doc, paragraph=None):
    """
    Return the right-tab stop position at the text area's right edge.
    Uses actual section settings from the document rather than hardcoded math.
    """
    section = doc.sections[0]
    right_tab_pos = section.page_width - section.left_margin - section.right_margin

    # Respect explicit right indent on the paragraph, if any.
    if paragraph and paragraph.paragraph_format and paragraph.paragraph_format.right_indent:
        right_tab_pos = right_tab_pos - paragraph.paragraph_format.right_indent

    return right_tab_pos


def add_hidden_text(run, text):
    """Add hidden text to a run for ATS parsing"""
    rPr = run._element.get_or_add_rPr()
    vanish = OxmlElement('w:vanish')
    rPr.append(vanish)
    run.text = text


def add_hyperlink_to_run(paragraph, url, text, font_size=None):
    """Add a hyperlink to a paragraph with optional font size"""
    # Decode HTML entities in text
    text = html.unescape(text)
    
    # Add the hyperlink relationship
    part = paragraph.part
    rId = part.relate_to(url, URL['link'], is_external=True)
    
    # Create the hyperlink
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(f"{{{URL['schema']}}}id", rId)
    
    # Create the run with formatting
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    color = OxmlElement('w:color')
    color.set(f"{{{URL['schema']}}}val", '0000FF')
    rPr.append(color)
    
    underline = OxmlElement('w:u')
    underline.set(f"{{{URL['schema']}}}val", 'single')
    rPr.append(underline)
    
    # Add font size if provided - ALWAYS set it to match surrounding text
    # If None, we should still try to get it from paragraph style
    if font_size is None:
        # Try to get from paragraph style as fallback
        if hasattr(paragraph, 'style') and paragraph.style:
            try:
                if hasattr(paragraph.style, 'font') and paragraph.style.font.size:
                    font_size = paragraph.style.font.size
            except (AttributeError, KeyError):
                pass
    
    if font_size is not None:
        sz = OxmlElement('w:sz')
        if hasattr(font_size, 'pt'):
            sz.set(f"{{{URL['schema']}}}val", str(int(font_size.pt * 2)))  # Word uses half-points
        else:
            sz.set(f"{{{URL['schema']}}}val", str(int(float(font_size) * 2)))
        rPr.append(sz)
    
    r.append(rPr)
    
    t = OxmlElement('w:t')
    t.set(f"{{{URL['ns']}}}space", 'preserve')
    t.text = text
    r.append(t)
    
    hyperlink.append(r)
    
    # Insert the hyperlink
    paragraph._element.append(hyperlink)
    
    return hyperlink


def process_text_with_hyperlinks(paragraph, element, strip_leading=False, css_styles=None, parent_css_props=None):
    """Process text that may contain hyperlinks and other inline elements"""
    if not hasattr(element, 'contents'):
        # It's a NavigableString - just add the text with spaces preserved
        text = element
        if strip_leading:
            text = text.lstrip()
        paragraph.add_run(text)  # Add as-is to preserve spaces
        return
    
    # If the element itself is a link (has href attribute), handle it directly
    if hasattr(element, 'get') and element.get('href'):
        url = element.get('href', '')
        text = element.string if element.string else ""
        if not text:
            text = element.get_text()  # Fallback to all text if no direct string
        if url and text:
            # Get font size from context - ALWAYS get from paragraph style first (most reliable)
            # Runs may not have explicit font.size if they inherit from paragraph style
            font_size = None
            
            # FIRST: Get font size from paragraph style (this is what text actually uses)
            if hasattr(paragraph, 'style') and paragraph.style:
                try:
                    if hasattr(paragraph.style, 'font') and paragraph.style.font.size:
                        font_size = paragraph.style.font.size
                except (AttributeError, KeyError):
                    pass
            
            # SECOND: Check if any run has explicit font size (overrides style)
            if paragraph.runs:
                for run in reversed(paragraph.runs):
                    if run.font.size:
                        font_size = run.font.size
                        break
            
            # THIRD: Try parent CSS props
            if font_size is None and parent_css_props and 'font-size' in parent_css_props:
                font_size = css_size_to_points(parent_css_props['font-size'])
            
            # FINAL fallback
            if font_size is None:
                font_size = Pt(BASE_FONT_SIZE)
            
            # Check if link has its own CSS classes that affect font size (override if present)
            if css_styles:
                link_classes = element.get('class', [])
                for link_class in link_classes:
                    if link_class in css_styles:
                        link_css_props = css_styles[link_class]
                        if 'font-size' in link_css_props:
                            font_size = css_size_to_points(link_css_props['font-size'])
            
            add_hyperlink_to_run(paragraph, url, text, font_size=font_size)
        return
    
    # Get effective CSS props for this element to determine font size context
    if css_styles:
        element_css_props = {}
        element_classes = element.get('class', [])
        for class_name in element_classes:
            if class_name in css_styles:
                element_css_props.update(css_styles[class_name])
        # Merge with parent props
        if parent_css_props:
            effective_css_props = parent_css_props.copy()
            effective_css_props.update(element_css_props)
        else:
            effective_css_props = element_css_props
    else:
        effective_css_props = parent_css_props or {}
    
    # Determine font size from context - prioritize CSS, then paragraph style, then existing runs
    font_size = None
    if 'font-size' in effective_css_props:
        font_size = css_size_to_points(effective_css_props['font-size'])
    elif parent_css_props and 'font-size' in parent_css_props:
        font_size = css_size_to_points(parent_css_props['font-size'])
    # If no font size from CSS, try paragraph style first (most reliable)
    if font_size is None and hasattr(paragraph, 'style') and paragraph.style:
        try:
            if hasattr(paragraph.style, 'font') and paragraph.style.font.size:
                font_size = paragraph.style.font.size
        except (AttributeError, KeyError):
            pass
    # Fallback to existing runs in paragraph
    if font_size is None and paragraph.runs:
        # Try to get font size from existing runs in the paragraph
        for run in paragraph.runs:
            if run.font.size:
                font_size = run.font.size
                break
    # Final fallback to base font size
    if font_size is None:
        font_size = Pt(BASE_FONT_SIZE)
    
    for idx, content in enumerate(element.contents):
        # Skip comments
        if isinstance(content, Comment):
            continue
        if isinstance(content, str):
            # Plain text string - preserve exactly as-is for spaces
            text = content
            if strip_leading and idx == 0:
                text = text.lstrip()
            paragraph.add_run(text)
        elif hasattr(content, 'get'):
            # It's a tag - check if it's a link (has href attribute)
            if content.get('href'):
                # Add hyperlink
                url = content.get('href', '')
                text = content.string if content.string else ""  # Get only direct text, not nested
                if not text:
                    text = content.get_text()  # Fallback to all text if no direct string
                if url and text:
                    # Get font size for this specific link element - ALWAYS get from paragraph style first
                    link_font_size = None
                    
                    # FIRST: Get font size from paragraph style (this is what text actually uses)
                    if hasattr(paragraph, 'style') and paragraph.style:
                        try:
                            if hasattr(paragraph.style, 'font') and paragraph.style.font.size:
                                link_font_size = paragraph.style.font.size
                        except (AttributeError, KeyError):
                            pass
                    
                    # SECOND: Check if any run has explicit font size (overrides style)
                    if paragraph.runs:
                        for run in reversed(paragraph.runs):
                            if run.font.size:
                                link_font_size = run.font.size
                                break
                    
                    # THIRD: Use the font_size we determined from context
                    if link_font_size is None:
                        link_font_size = font_size
                    
                    # Check if link has its own CSS classes that override font size
                    if css_styles:
                        link_classes = content.get('class', [])
                        link_css_props = effective_css_props.copy()
                        for link_class in link_classes:
                            if link_class in css_styles:
                                link_css_props.update(css_styles[link_class])
                        if 'font-size' in link_css_props:
                            link_font_size = css_size_to_points(link_css_props['font-size'])
                    
                    add_hyperlink_to_run(paragraph, url, text, font_size=link_font_size)
                # Don't recurse into <a> - we already got the text
            elif content.name == 'span':
                # Check if this is a separator element
                content_classes = content.get('class', [])
                is_separator = any(any(pattern in str(c).lower() for pattern in SEPARATOR_CLASS_PATTERNS) for c in content_classes)
                if is_separator:
                    # Replace separator with separator text
                    paragraph.add_run(SEPARATOR_REPLACEMENT)
                else:
                    # Regular span - process normally, pass CSS props
                    process_text_with_hyperlinks(paragraph, content, False, css_styles, parent_css_props=effective_css_props)
            elif content.name == 'i' or content.name == 'em':
                # Add italic text
                run = paragraph.add_run(content.get_text())
                run.font.italic = True
                if font_size:
                    run.font.size = font_size
            elif content.name == 'b' or content.name == 'strong':
                # Add bold text
                run = paragraph.add_run(content.get_text())
                run.font.bold = True
                if font_size:
                    run.font.size = font_size
            else:
                # Recursively process nested elements, pass CSS props
                process_text_with_hyperlinks(paragraph, content, False, css_styles, parent_css_props=effective_css_props)


def parse_inline_style(style_attr):
    """Parse inline style attribute (e.g., 'font-weight: normal; font-style: italic')"""
    props = {}
    if not style_attr:
        return props
    for prop in style_attr.split(';'):
        prop = prop.strip()
        if ':' in prop:
            key, val = prop.split(':', 1)
            props[key.strip()] = val.strip()
    return props


def parse_css_variables(css):
    """Parse CSS custom properties (variables) from :root"""
    variables = {}
    # Match :root { ... }
    root_match = re.search(r':root\s*\{([^}]+)\}', css, re.DOTALL)
    if root_match:
        root_content = root_match.group(1)
        # Parse --variable-name: value;
        for match in re.finditer(r'--([^:]+):\s*([^;]+);', root_content):
            var_name = match.group(1).strip()
            var_value = match.group(2).strip()
            variables[var_name] = var_value
    return variables


def resolve_css_variable(value, variables):
    """Resolve CSS var() references to their actual values"""
    if not value or 'var(' not in value:
        return value
    
    # Match var(--variable-name) or var(--variable-name, fallback)
    var_pattern = r'var\(\s*--([^,)]+)\s*(?:,\s*([^)]+))?\s*\)'
    
    def replacer(match):
        var_name = match.group(1).strip()
        fallback = match.group(2).strip() if match.group(2) else None
        
        # Look up variable
        if var_name in variables:
            return variables[var_name]
        elif fallback:
            return fallback
        else:
            return match.group(0)  # Return original if not found
    
    return re.sub(var_pattern, replacer, value)


def parse_css(soup):
    """Parse CSS styles - handles comma-separated selectors like .patent, .degree, .position"""
    styles = {}
    style_tag = soup.find('style')
    if style_tag and style_tag.string:
        css = style_tag.string
        
        # First, parse CSS variables from :root
        css_variables = parse_css_variables(css)
        # Match CSS rules - find selector block { properties }
        for match in re.finditer(r'([^{]+)\{([^}]+)\}', css):
            selector_part = match.group(1).strip()
            props_text = match.group(2)

            # Parse properties and resolve CSS variables
            props = {}
            for prop in props_text.split(';'):
                prop = prop.strip()
                if ':' in prop:
                    key, val = prop.split(':', 1)
                    # Resolve any var() references in the value
                    resolved_val = resolve_css_variable(val.strip(), css_variables)
                    props[key.strip()] = resolved_val

            # Split by comma to handle multiple selectors
            selectors = [s.strip() for s in selector_part.split(',')]
            
            for selector in selectors:
                # Skip pseudo-elements (::before, ::after) and pseudo-classes (:hover, :link, :visited)
                # These don't correspond to actual HTML elements/classes
                if '::' in selector or ':' in selector:
                    continue
                
                # Check if this is a descendant selector (e.g., ".secondary .position")
                descendant_match = re.match(r'\.(\w+)\s+\.(\w+)', selector)
                if descendant_match:
                    # This is a descendant selector - create a flattened style name
                    parent_class = descendant_match.group(1)
                    child_class = descendant_match.group(2)
                    # Create a combined style name like "secondary_position"
                    combined_name = f"{parent_class}_{child_class}"
                    if combined_name in styles:
                        styles[combined_name].update(props)
                    else:
                        styles[combined_name] = props.copy()
                else:
                    # Regular class selector - extract class name
                    class_match = re.search(r'\.(\w+)', selector)
                    if class_match:
                        class_name = class_match.group(1)
                        # Only apply to base class if it's not part of a descendant selector in this rule
                        # (we'll handle descendant selectors separately above)
                        if class_name in styles:
                            styles[class_name].update(props)
                        else:
                            styles[class_name] = props.copy()
    return styles


def get_class_importance(class_name, css_props):
    """Calculate importance/hierarchy level from CSS properties (1 = most important)"""
    # Structural/container classes - identified by having display:grid or similar layout properties
    if 'display' in css_props and ('grid' in css_props['display'] or 'flex' in css_props['display']):
        return 10
    
    # Calculate importance from font-size (larger = more important)
    importance = 5  # Default middle importance
    
    if 'font-size' in css_props:
        font_size = css_props['font-size']
        if '%' in font_size:
            percent = float(font_size.strip('%'))
            if percent >= 140:
                importance = 1  # Very large (headers)
            elif percent >= 120:
                importance = 2  # Large (sub-headers)
            elif percent >= 100:
                importance = 3  # Normal to large
            elif percent >= 90:
                importance = 4  # Slightly smaller
            else:
                importance = 5  # Small
        elif 'em' in font_size or 'rem' in font_size:
            multiplier = float(re.sub(r'[^\d.]', '', font_size) or '1')
            if multiplier >= 1.4:
                importance = 1
            elif multiplier >= 1.2:
                importance = 2
            elif multiplier >= 1.0:
                importance = 3
            elif multiplier >= 0.9:
                importance = 4
            else:
                importance = 5
    
    # Boost importance if bold
    if 'font-weight' in css_props:
        weight = css_props['font-weight']
        if 'bold' in weight.lower() or 'bolder' in weight.lower() or weight in ['600', '700', '800', '900']:
            importance = max(1, importance - 1)
    
    # Boost importance if small-caps (usually headers)
    if 'font-variant' in css_props and 'small-caps' in css_props['font-variant']:
        importance = max(1, importance - 1)
    
    # Reduce importance if italic (often metadata/details)
    if 'font-style' in css_props and 'italic' in css_props['font-style']:
        importance = min(10, importance + 1)
    
    # Reduce importance if small color (grey/dim - often metadata)
    if 'color' in css_props:
        color = css_props['color'].lower()
        if 'grey' in color or 'gray' in color or 'dim' in color:
            importance = min(10, importance + 1)
    
    return importance


def css_size_to_points(css_size, base_size=BASE_FONT_SIZE):
    """Convert CSS font-size to Word points"""
    if not css_size:
        return None
    
    css_size = css_size.strip()
    
    # Percentage
    if '%' in css_size:
        percent = float(css_size.strip('%')) / 100
        return Pt(base_size * percent)
    
    # em/rem (treat as multiplier of base)
    if 'em' in css_size or 'rem' in css_size:
        multiplier = float(re.sub(r'[^\d.]', '', css_size))
        return Pt(base_size * multiplier)
    
    # Points already
    if 'pt' in css_size:
        return Pt(float(css_size.strip('pt')))
    
    # Pixels (rough conversion: 1px ≈ 0.75pt)
    if 'px' in css_size:
        return Pt(float(css_size.strip('px')) * 0.75)
    
    # Try to parse as number (assume points)
    try:
        return Pt(float(css_size))
    except:
        return None


def css_margin_to_inches(css_margin, base_size=BASE_FONT_SIZE):
    """Convert CSS margin value to Word spacing in inches (or points)"""
    if not css_margin:
        return None

    css_margin = css_margin.strip()

    # Zero means no spacing - return Pt(0) explicitly
    css_margin_clean = css_margin.strip().lower()
    if css_margin_clean == '0' or css_margin_clean == '0lh' or css_margin_clean == '0ch' or css_margin_clean == '0px' or css_margin_clean == '0pt':
        return Pt(0)

    # lh units (line-height - vertical unit)
    # 1lh equals the computed line-height, which is typically 1.0-1.2 times font size
    # For simplicity, we approximate 1lh ≈ 1.0 * font_size
    if 'lh' in css_margin:
        lh_value = float(re.sub(r'[^\d.-]', '', css_margin) or '0')
        # Convert lh to points: 1lh ≈ 1.0 * font_size
        pt_value = lh_value * base_size
        return Pt(pt_value)
    
    # ch units (character width - width of '0' character in current font)
    # 1ch is approximately 0.5-0.6em width, but for vertical spacing we use a smaller multiplier
    # For vertical margins/padding, 1ch typically renders as about 0.3-0.4em in practice
    if 'ch' in css_margin:
        ch_value = float(re.sub(r'[^\d.-]', '', css_margin) or '0')
        # Convert ch to points: 1ch ≈ 0.35 * font_size for vertical spacing
        pt_value = ch_value * base_size * 0.35
        # Word spacing can use Points directly, which is more accurate
        return Pt(pt_value)
    
    # em/rem units (relative to font size)
    if 'em' in css_margin or 'rem' in css_margin:
        em_value = float(re.sub(r'[^\d.-]', '', css_margin) or '0')
        # Convert to points then to inches (72 points = 1 inch)
        pt_value = base_size * em_value
        return Inches(pt_value / 72)
    
    # Points
    if 'pt' in css_margin:
        pt_value = float(css_margin.strip('pt'))
        return Inches(pt_value / 72)
    
    # Inches
    if 'in' in css_margin:
        return Inches(float(css_margin.strip('in')))
    
    # Pixels (1px ≈ 0.75pt ≈ 0.0104 inches)
    if 'px' in css_margin:
        px_value = float(re.sub(r'[^\d.-]', '', css_margin) or '0')
        return Inches(px_value * 0.0104)
    
    # Try to parse as number (assume points)
    try:
        pt_value = float(css_margin)
        return Inches(pt_value / 72)
    except:
        return None


def derive_style_name(class_name, css_props, importance_level):
    """Derive a Word style name from class name and CSS properties"""
    # Use the class name directly, converting underscores to spaces and capitalizing appropriately
    style_name = class_name.replace('_', ' ').replace('-', ' ')
    
    # Capitalize words appropriately
    words = style_name.split()
    style_name = ' '.join(word.capitalize() for word in words)
    
    return style_name


def create_word_style_from_css(doc, class_name, css_props, importance_level, is_paragraph=True):
    """Create a Word style from CSS properties
    
    Args:
        doc: Word document
        class_name: CSS class name
        css_props: CSS properties dict
        importance_level: Style importance level
        is_paragraph: True for paragraph style (div classes), False for character style (span classes)
    """
    style_name = derive_style_name(class_name, css_props, importance_level)
    
    if is_paragraph:
        style_type = WD_STYLE_TYPE.PARAGRAPH
    else:
        style_type = WD_STYLE_TYPE.CHARACTER
    
    try:
        # Check if style already exists
        existing_style = doc.styles[style_name]
        # Check if existing style matches what we want
        if (is_paragraph and hasattr(existing_style, 'paragraph_format')) or \
           (not is_paragraph and not hasattr(existing_style, 'paragraph_format')):
            style = existing_style
        else:
            # Style exists but wrong type - create with different name
            if is_paragraph:
                style_name = f"{style_name} Para"
            else:
                style_name = f"{style_name} Char"
            style = doc.styles.add_style(style_name, style_type)
    except KeyError:
        # Create new style
        style = doc.styles.add_style(style_name, style_type)
    
    # Set paragraph format (only if it's a paragraph style)
    if not hasattr(style, 'paragraph_format'):
        # Fallback: can't set paragraph format on character style
        pf = None
    else:
        pf = style.paragraph_format
    
    # Get font size for line-height calculations
    base_size = BASE_FONT_SIZE  # Default
    if 'font-size' in css_props:
        font_size = css_size_to_points(css_props['font-size'])
        if font_size:
            style.font.size = font_size
            # Extract point value for line-height calculations
            base_size = font_size.pt if hasattr(font_size, 'pt') else BASE_FONT_SIZE
    
    # Font weight
    if 'font-weight' in css_props:
        weight = css_props['font-weight']
        if 'bold' in weight.lower() or 'bolder' in weight.lower() or weight in ['600', '700', '800', '900']:
            style.font.bold = True
    
    # Font style
    if 'font-style' in css_props:
        if 'italic' in css_props['font-style']:
            style.font.italic = True
    
    # Color
    if 'color' in css_props:
        color = css_props['color']
        if 'dimgrey' in color or 'dimgray' in color:
            style.font.color.rgb = RGBColor(105, 105, 105)
        elif 'grey' in color.lower() or 'gray' in color.lower():
            style.font.color.rgb = RGBColor(128, 128, 128)
    
    # Small caps (font-variant)
    if 'font-variant' in css_props and 'small-caps' in css_props['font-variant']:
        style.font.small_caps = True
    
    # Spacing derived from CSS margins (only for paragraph styles)
    if pf is not None:
        # margin-top maps to space_before (space above the paragraph)
        if 'margin-top' in css_props:
            space_before = css_margin_to_inches(css_props['margin-top'])
            if space_before is not None:
                pf.space_before = space_before
        else:
            # If margin-top not specified, default to no space before
            pf.space_before = Pt(0)
        
        # margin-bottom maps to space_after (space below the paragraph)
        if 'margin-bottom' in css_props:
            space_after = css_margin_to_inches(css_props['margin-bottom'])
            if space_after is not None:
                pf.space_after = space_after
        else:
            # If margin-bottom not specified, explicitly set to 0 (no space after)
            # This prevents Word from applying default spacing
            pf.space_after = Pt(0)
        
        # Also check padding-top and padding-bottom as fallbacks
        # (padding is less common for vertical spacing in CSS but some use it)
        if 'padding-top' in css_props and 'margin-top' not in css_props:
            space_before = css_margin_to_inches(css_props['padding-top'])
            if space_before is not None:
                pf.space_before = space_before
        
        if 'padding-bottom' in css_props and 'margin-bottom' not in css_props:
            space_after = css_margin_to_inches(css_props['padding-bottom'])
            if space_after is not None:
                pf.space_after = space_after
        
        # Line spacing from CSS line-height property
        if 'line-height' in css_props:
            line_height = css_props['line-height'].strip()
            
            # Percentage (e.g., "100%", "150%")
            if '%' in line_height:
                percent = float(line_height.strip('%')) / 100
                if percent == 1.0:
                    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    pf.line_spacing = None
                elif percent == 1.5:
                    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                    pf.line_spacing = None
                elif percent == 2.0:
                    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                    pf.line_spacing = None
                else:
                    # Custom multiple
                    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    pf.line_spacing = percent
            
            # Unitless number (e.g., "1", "1.5", "2")
            elif line_height.replace('.', '').isdigit():
                value = float(line_height)
                if value == 1.0:
                    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    pf.line_spacing = None
                elif value == 1.5:
                    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                    pf.line_spacing = None
                elif value == 2.0:
                    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                    pf.line_spacing = None
                else:
                    # Custom multiple
                    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    pf.line_spacing = value
            
            # Absolute length (e.g., "12pt", "1em")
            elif 'pt' in line_height or 'px' in line_height or 'em' in line_height or 'rem' in line_height:
                # Convert to points - need font size for em/rem
                font_size_pt = base_size
                if 'font-size' in css_props:
                    fs = css_size_to_points(css_props['font-size'])
                    if fs and hasattr(fs, 'pt'):
                        font_size_pt = fs.pt
                
                if 'pt' in line_height:
                    pt_value = float(line_height.strip('pt'))
                elif 'px' in line_height:
                    px_value = float(re.sub(r'[^\d.]', '', line_height) or '0')
                    pt_value = px_value * 0.75
                elif 'em' in line_height or 'rem' in line_height:
                    em_value = float(re.sub(r'[^\d.]', '', line_height) or '0')
                    pt_value = font_size_pt * em_value
                
                # Use AT_LEAST for absolute line heights
                pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                pf.line_spacing = Pt(pt_value)
        
        # Fallback: if no line-height specified, use single spacing as default
        elif 'line-height' not in css_props:
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            pf.line_spacing = None
        
        # Alignment
        if 'text-align' in css_props:
            align = css_props['text-align']
            if 'center' in align:
                pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif 'right' in align:
                pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif 'left' in align:
                pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Left indent from margin-left or padding-left
        if 'margin-left' in css_props:
            left_indent = css_margin_to_inches(css_props['margin-left'])
            if left_indent is not None:
                pf.left_indent = left_indent
        elif 'padding-left' in css_props:
            left_indent = css_margin_to_inches(css_props['padding-left'])
            if left_indent is not None:
                pf.left_indent = left_indent
    
    return style


def create_all_word_styles(doc, css_styles, soup):
    """Create Word styles for all relevant CSS classes
    
    Args:
        doc: Word document
        css_styles: Dict of CSS class names to CSS properties
        soup: BeautifulSoup parsed HTML to determine div vs span classes
    """
    created_styles = {}
    skipped_classes = []
    
    # Determine which classes are used on divs vs spans
    div_classes = set()
    span_classes = set()
    
    for div in soup.find_all('div'):
        classes = div.get('class', [])
        div_classes.update(classes)
    
    for span in soup.find_all('span'):
        classes = span.get('class', [])
        span_classes.update(classes)
    
    # Create styles for all classes, including structural containers
    # They can have margin-top and other spacing that should be applied
    for class_name, css_props in css_styles.items():
        # Skip classes that have no meaningful CSS properties (likely decorators/separators)
        # These would have been replaced in HTML preprocessing anyway
        if not css_props or (len(css_props) == 1 and 'display' in css_props):
            skipped_classes.append(f"{class_name} (decorator/separator)")
            continue
        
        # Determine if this should be a paragraph style (div) or character style (span)
        is_span_class = class_name in span_classes and class_name not in div_classes
        is_div_class = class_name in div_classes
        
        # Handle synthetic classes like "secondary_position" - if it ends with a known span class,
        # treat it as a character style
        is_synthetic_span = False
        if '_' in class_name:
            # Check if the suffix matches a known span class
            suffix = class_name.split('_')[-1]
            if suffix in span_classes and suffix not in div_classes:
                is_synthetic_span = True
        
        # Default to paragraph style if used on divs, character style if only on spans or synthetic span
        style_type_is_paragraph = is_div_class or (not is_span_class and not is_div_class and not is_synthetic_span)
        
        # Calculate importance dynamically from CSS
        importance = get_class_importance(class_name, css_props)
        word_style = create_word_style_from_css(doc, class_name, css_props, importance, is_paragraph=style_type_is_paragraph)
        created_styles[class_name] = word_style.name
        
        # Note layout containers but still create styles for them
        if 'display' in css_props and ('grid' in css_props['display'] or 'flex' in css_props['display']):
            skipped_classes.append(f"{class_name} (layout container, but style created for spacing)")
    
    # Print detailed mapping with all requested properties
    # First pass: collect all values to determine column widths
    all_data = []
    for css_class, word_style_name in sorted(created_styles.items()):
        css_props = css_styles[css_class]
        
        # Extract CSS values
        css_font_size = css_props.get('font-size', '100%')
        css_font_weight = css_props.get('font-weight', 'normal')
        css_font_style = css_props.get('font-style', 'normal')
        css_line_height = css_props.get('line-height', 'normal')
        css_margin_top = css_props.get('margin-top', css_props.get('padding-top', '0'))
        css_margin_bottom = css_props.get('margin-bottom', css_props.get('padding-bottom', '0'))
        css_margin_left = css_props.get('margin-left', css_props.get('padding-left', '0'))
        
        # Format CSS style (bold/italic/normal)
        css_style = "normal"
        if 'bold' in css_font_weight.lower() or 'bolder' in css_font_weight.lower() or css_font_weight in ['600', '700', '800', '900']:
            if 'italic' in css_font_style.lower():
                css_style = "bold+italic"
            else:
                css_style = "bold"
        elif 'italic' in css_font_style.lower():
            css_style = "italic"
        
        # Get Word style and extract values
        try:
            word_style = doc.styles[word_style_name]
            pf = word_style.paragraph_format if hasattr(word_style, 'paragraph_format') else None
            
            # Word font size
            word_font_size = f"{BASE_FONT_SIZE}pt"
            if hasattr(word_style, 'font') and word_style.font.size:
                if hasattr(word_style.font.size, 'pt'):
                    word_font_size = f"{word_style.font.size.pt:.1f}pt"
                else:
                    word_font_size = str(word_style.font.size)
            
            # Word style (bold/italic/normal)
            word_style_desc = "normal"
            if hasattr(word_style, 'font'):
                is_bold = word_style.font.bold if hasattr(word_style.font, 'bold') else False
                is_italic = word_style.font.italic if hasattr(word_style.font, 'italic') else False
                if is_bold and is_italic:
                    word_style_desc = "bold+italic"
                elif is_bold:
                    word_style_desc = "bold"
                elif is_italic:
                    word_style_desc = "italic"
            
            # Word line height/spacing
            word_line_height = "1.0"
            if pf:
                if pf.line_spacing_rule == WD_LINE_SPACING.SINGLE:
                    word_line_height = "1.0"
                elif pf.line_spacing_rule == WD_LINE_SPACING.ONE_POINT_FIVE:
                    word_line_height = "1.5"
                elif pf.line_spacing_rule == WD_LINE_SPACING.DOUBLE:
                    word_line_height = "2.0"
                elif pf.line_spacing_rule == WD_LINE_SPACING.MULTIPLE:
                    if pf.line_spacing:
                        if hasattr(pf.line_spacing, '__float__'):
                            word_line_height = f"{float(pf.line_spacing):.2f}"
                        else:
                            word_line_height = str(pf.line_spacing)
                    else:
                        word_line_height = "multiple"
                elif pf.line_spacing_rule == WD_LINE_SPACING.AT_LEAST:
                    if pf.line_spacing:
                        if hasattr(pf.line_spacing, 'pt'):
                            word_line_height = f"{pf.line_spacing.pt:.1f}pt"
                        else:
                            word_line_height = f"{pf.line_spacing}pt"
                    else:
                        word_line_height = "at least"
            
            # Word leading vertical space (space_before)
            word_lead_space = "0"
            if pf and pf.space_before:
                if hasattr(pf.space_before, 'pt'):
                    word_lead_space = f"{pf.space_before.pt:.1f}pt"
                elif hasattr(pf.space_before, '__float__'):
                    word_lead_space = f"{float(pf.space_before):.1f}pt"
                else:
                    word_lead_space = str(pf.space_before)
            
            # Word trailing vertical space (space_after)
            word_trail_space = "0"
            if pf and pf.space_after:
                if hasattr(pf.space_after, 'pt'):
                    word_trail_space = f"{pf.space_after.pt:.1f}pt"
                elif hasattr(pf.space_after, '__float__'):
                    word_trail_space = f"{float(pf.space_after):.1f}pt"
                else:
                    word_trail_space = str(pf.space_after)
            
            # Word left indent
            word_left_indent = "0"
            if pf and pf.left_indent:
                if hasattr(pf.left_indent, 'pt'):
                    word_left_indent = f"{pf.left_indent.pt:.1f}pt"
                elif hasattr(pf.left_indent, 'inches'):
                    word_left_indent = f"{pf.left_indent.inches:.2f}in"
                elif hasattr(pf.left_indent, '__float__'):
                    word_left_indent = f"{float(pf.left_indent):.1f}pt"
                else:
                    word_left_indent = str(pf.left_indent)
        except Exception as e:
            word_font_size = "?"
            word_style_desc = "?"
            word_line_height = "?"
            word_lead_space = "?"
            word_trail_space = "?"
            word_left_indent = "?"
        
        all_data.append({
            'css_class': css_class,
            'word_style_name': word_style_name,
            'css_font_size': css_font_size,
            'css_style': css_style,
            'css_line_height': css_line_height,
            'css_margin_top': css_margin_top,
            'css_margin_bottom': css_margin_bottom,
            'css_margin_left': css_margin_left,
            'word_font_size': word_font_size,
            'word_style_desc': word_style_desc,
            'word_line_height': word_line_height,
            'word_lead_space': word_lead_space,
            'word_trail_space': word_trail_space,
            'word_left_indent': word_left_indent
        })
    
    # Calculate column widths based on longest values (including header labels)
    width_css_class = max([len('CSS Class')] + [len(d['css_class']) for d in all_data]) + 2
    width_word_style = max([len('Word Style')] + [len(d['word_style_name']) for d in all_data]) + 2
    width_css_size = max([len('size')] + [len(d['css_font_size']) for d in all_data]) + 2
    width_css_style = max([len('style')] + [len(d['css_style']) for d in all_data]) + 2
    width_css_line_h = max([len('line-h')] + [len(d['css_line_height']) for d in all_data]) + 2
    width_css_space_top = max([len('spc-top')] + [len(d['css_margin_top']) for d in all_data]) + 2
    width_css_space_bot = max([len('spc-bot')] + [len(d['css_margin_bottom']) for d in all_data]) + 2
    width_css_indent = max([len('indent')] + [len(d['css_margin_left']) for d in all_data]) + 2
    width_word_size = max([len('size')] + [len(d['word_font_size']) for d in all_data]) + 2
    width_word_style_desc = max([len('style')] + [len(d['word_style_desc']) for d in all_data]) + 2
    width_word_line_h = max([len('line-h')] + [len(d['word_line_height']) for d in all_data]) + 2
    width_word_space_top = max([len('spc-top')] + [len(d['word_lead_space']) for d in all_data]) + 2
    width_word_space_bot = max([len('spc-bot')] + [len(d['word_trail_space']) for d in all_data]) + 2
    width_word_indent = max([len('indent')] + [len(d['word_left_indent']) for d in all_data]) + 2
    
    # Print header - align header labels with data column widths
    print()
    css_header_part = f"[{'size':<{width_css_size-1}}, {'style':<{width_css_style-1}}, {'line-h':<{width_css_line_h-1}}, {'spc-top':<{width_css_space_top-1}}, {'spc-bot':<{width_css_space_bot-1}}, {'indent':<{width_css_indent-1}}]"
    word_header_part = f"[{'size':<{width_word_size-1}}, {'style':<{width_word_style_desc-1}}, {'line-h':<{width_word_line_h-1}}, {'spc-top':<{width_word_space_top-1}}, {'spc-bot':<{width_word_space_bot-1}}, {'indent':<{width_word_indent-1}}]"
    header = f"{'CSS Class':<{width_css_class}} {css_header_part} → {'Word Style':<{width_word_style}} {word_header_part}"
    print(header)
    
    # Separator line - calculate total width
    separator_css = width_css_class + len(css_header_part) + 3  # +3 for spaces around bracket/arrow
    separator_word = width_word_style + len(word_header_part) + 3
    print(f"{'-'*separator_css} → {'-'*separator_word}")
    
    # Print data rows
    for d in all_data:
        css_part = f"[{d['css_font_size']:<{width_css_size-1}}, {d['css_style']:<{width_css_style-1}}, {d['css_line_height']:<{width_css_line_h-1}}, {d['css_margin_top']:<{width_css_space_top-1}}, {d['css_margin_bottom']:<{width_css_space_bot-1}}, {d['css_margin_left']:<{width_css_indent-1}}]"
        word_part = f"[{d['word_font_size']:<{width_word_size-1}}, {d['word_style_desc']:<{width_word_style_desc-1}}, {d['word_line_height']:<{width_word_line_h-1}}, {d['word_lead_space']:<{width_word_space_top-1}}, {d['word_trail_space']:<{width_word_space_bot-1}}, {d['word_left_indent']:<{width_word_indent-1}}]"
        print(f"{d['css_class']:<{width_css_class}} {css_part} → {d['word_style_name']:<{width_word_style}} {word_part}")
    
    if skipped_classes:
        print(f"\nSkipped classes (structural/container):")
        for skipped in skipped_classes:
            print(f"  - {skipped}")
    
    return created_styles


def is_inside_parent(element, parent_class):
    """Check if an element is inside a parent with the given class"""
    parent = element.parent
    while parent:
        if hasattr(parent, 'get') and parent.get('class'):
            classes = parent.get('class', [])
            if parent_class in classes:
                return True
        parent = getattr(parent, 'parent', None)
    return False


def resolve_descendant_style(element, class_name, css_styles):
    """Resolve if a class should use a descendant style variant
    
    Checks if the element's class has a descendant style variant (e.g., 'parent_child')
    by checking if the element is inside a parent with the parent class.
    
    Args:
        element: BeautifulSoup element
        class_name: The CSS class name to check
        css_styles: Dictionary of CSS styles (may contain descendant styles like 'parent_child')
    
    Returns:
        The effective class name to use (either descendant variant or original class_name)
    """
    # Check all styles for descendant patterns (those with underscores)
    for style_name in css_styles.keys():
        if '_' in style_name:
            # Split into parent and child parts
            parts = style_name.split('_')
            if len(parts) >= 2:
                # Check if the last part matches our class name
                child_part = parts[-1]
                if child_part == class_name:
                    # Check if element is inside a parent with the parent class(es)
                    # Handle multiple parent levels (e.g., "grandparent_parent_child")
                    parent_classes = parts[:-1]
                    for parent_class in parent_classes:
                        if is_inside_parent(element, parent_class):
                            # Found a matching descendant style
                            return style_name
    # No descendant style found, use original class name
    return class_name


def apply_style(run, class_name, styles):
    """Apply CSS styles to a run - allows mixing bold/normal text within same paragraph"""
    if not run or class_name not in styles:
        return
    
    style = styles[class_name]
    
    # Font size - apply at run level for fine-grained control
    if 'font-size' in style and '%' in style['font-size']:
        percent = float(style['font-size'].strip('%')) / 100
        run.font.size = Pt(BASE_FONT_SIZE * percent)
    
    # Font weight - apply at run level to allow mixed bold/normal in same paragraph
    if 'font-weight' in style and ('bold' in style['font-weight'] or 'bolder' in style['font-weight']):
        run.font.bold = True
    
    if 'font-style' in style and 'italic' in style['font-style']:
        # Italic can be at run level for emphasis within styled text
        run.font.italic = True
    
    # Color can vary at run level (e.g., dimgrey for dates)
    if 'color' in style and 'dimgrey' in style['color']:
        run.font.color.rgb = RGBColor(105, 105, 105)
    elif 'color' in style and ('grey' in style['color'].lower() or 'gray' in style['color'].lower()):
        run.font.color.rgb = RGBColor(128, 128, 128)


def create_bulleted_style(doc):
    """Create a custom bulleted style with single line spacing and no paragraph spacing"""
    try:
        # Try to get existing style
        style = doc.styles['bulleted']
    except KeyError:
        # Create new style based on List Bullet
        list_bullet = doc.styles['List Bullet']
        style = doc.styles.add_style('bulleted', WD_STYLE_TYPE.PARAGRAPH)
        
        # Copy numbering format from List Bullet to get bullets
        list_pPr = list_bullet.element.pPr
        if list_pPr is not None:
            list_numPr = list_pPr.numPr
            if list_numPr is not None:
                # Get or create pPr for our style
                style_pPr = style.element.get_or_add_pPr()
                # Create new numPr element and copy the XML
                new_numPr = OxmlElement('w:numPr')
                # Copy all child elements
                for child in list_numPr:
                    new_numPr.append(child)
                # Insert the numPr element into pPr
                style_pPr.insert(0, new_numPr)
    
    # Set paragraph format
    pf = style.paragraph_format
    pf.left_indent = BULLET_LEFT_INDENT
    pf.first_line_indent = BULLET_FIRST_LINE_INDENT
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    # For SINGLE spacing, line_spacing should be None, not 1.0
    pf.line_spacing = None
    
    return style


def convert_html_to_docx(html_file, output_file=None):
    """Convert HTML to DOCX with proper formatting"""
    
    with open(html_file, 'r', encoding='utf-8') as f:
        html_content = f.read()
    
    # Replace all separator elements with separator text before parsing
    # This is the single generic spot for this replacement - happens before any parsing
    # Handle multiline spans and various quote styles
    for separator_pattern in SEPARATOR_CLASS_PATTERNS:
        # Match opening and closing tags: <span class="vertical-bar"></span>
        html_content = re.sub(
            rf'<span[^>]*\s+class\s*=\s*["\']{re.escape(separator_pattern)}["\'][^>]*>\s*</span>',
            SEPARATOR_REPLACEMENT,
            html_content,
            flags=re.IGNORECASE | re.DOTALL
        )
        # Match self-closing tags: <span class="vertical-bar"/>
        html_content = re.sub(
            rf'<span[^>]*\s+class\s*=\s*["\']{re.escape(separator_pattern)}["\'][^>]*/>',
            SEPARATOR_REPLACEMENT,
            html_content,
            flags=re.IGNORECASE | re.DOTALL
        )
    
    soup = BeautifulSoup(html_content, 'html.parser')
    css_styles = parse_css(soup)
    
    doc = Document()
    
    # Create all Word styles from CSS
    word_styles = create_all_word_styles(doc, css_styles, soup)
    print(f"Created {len(word_styles)} Word styles from CSS classes")
    
    # Create custom bulleted style
    create_bulleted_style(doc)
    
    # Set page dimensions using constants
    section = doc.sections[0]
    section.page_width = Inches(PAGE_WIDTH)
    section.page_height = Inches(PAGE_HEIGHT)
    section.left_margin = Inches(PAGE_LEFT_MARGIN)
    section.right_margin = Inches(PAGE_RIGHT_MARGIN)
    section.top_margin = Inches(PAGE_TOP_MARGIN)
    section.bottom_margin = Inches(PAGE_BOTTOM_MARGIN)
    section.gutter = Inches(0)  # No gutter needed when margins are uniform
    
    # Track content height for pagination (in points)
    # Usable page height = 11" - 0.6" top - 0.6" bottom = 9.8" = 705.6 points
    current_page_height = 0
    usable_page_height_pt = USABLE_PAGE_HEIGHT * 72  # Convert inches to points
    
    # Find all top-level divs in body
    body = soup.find('body') or soup
    
    def estimate_paragraph_height(p):
        """Estimate paragraph height in points"""
        pf = p.paragraph_format
        height = 0
        
        # Space before
        if pf.space_before:
            if hasattr(pf.space_before, 'pt'):
                height += pf.space_before.pt
            else:
                height += float(str(pf.space_before).replace('Pt', '').replace('(', '').replace(')', '') or '0')
        
        # Line height * number of lines (estimate)
        font_size = BASE_FONT_SIZE
        if p.runs and p.runs[0].font.size:
            if hasattr(p.runs[0].font.size, 'pt'):
                font_size = p.runs[0].font.size.pt
        
        line_height_multiplier = 1.0
        if pf.line_spacing_rule == WD_LINE_SPACING.SINGLE:
            line_height_multiplier = 1.0
        elif pf.line_spacing_rule == WD_LINE_SPACING.ONE_POINT_FIVE:
            line_height_multiplier = 1.5
        elif pf.line_spacing_rule == WD_LINE_SPACING.DOUBLE:
            line_height_multiplier = 2.0
        elif pf.line_spacing and hasattr(pf.line_spacing, '__float__'):
            line_height_multiplier = float(pf.line_spacing)
        
        line_height = font_size * line_height_multiplier
        
        # Estimate number of lines (rough - assume 7.3" text width / average char width)
        text = p.text
        avg_char_width = font_size * 0.6  # Rough estimate
        chars_per_line = (7.3 * 72) / avg_char_width  # Text width in points
        num_lines = max(1, len(text) / chars_per_line) if text else 1
        
        height += line_height * num_lines
        
        # Space after
        if pf.space_after:
            if hasattr(pf.space_after, 'pt'):
                height += pf.space_after.pt
            else:
                height += float(str(pf.space_after).replace('Pt', '').replace('(', '').replace(')', '') or '0')
        
        return height
    
    def check_and_insert_page_break(p, current_height):
        """Check if we need a page break and insert one if needed"""
        para_height = estimate_paragraph_height(p)
        if current_height + para_height > usable_page_height_pt:
            # Insert page break before this paragraph
            run = p.runs[0] if p.runs else p.add_run()
            run.add_break(WD_BREAK.PAGE)
            return 0  # Reset height counter
        return current_height + para_height
    
    def get_effective_css_props(element, css_styles):
        """Get effective CSS properties for an element, considering all its classes and descendant selectors"""
        classes = element.get('class', [])
        effective_props = {}
        
        # Merge properties from all classes
        for class_name in classes:
            if class_name in css_styles:
                effective_props.update(css_styles[class_name])
        
        # Check for descendant style variants
        for class_name in classes:
            effective_class = resolve_descendant_style(element, class_name, css_styles)
            if effective_class != class_name and effective_class in css_styles:
                # Descendant style found - merge its properties (they override base)
                effective_props.update(css_styles[effective_class])
        
        # Merge inline styles (highest priority)
        inline_style = parse_inline_style(element.get('style', ''))
        effective_props.update(inline_style)
        
        return effective_props
    
    def process_element(element, parent_paragraph=None, parent_css_props=None):
        """Generic element processor - processes any HTML element based on CSS and structure"""
        nonlocal current_page_height
        if not element or not hasattr(element, 'name'):
            return current_page_height
        
        tag_name = element.name
        classes = element.get('class', [])
        css_props = get_effective_css_props(element, css_styles)
        
        # Inherit text-align from parent if not explicitly set
        if parent_css_props and 'text-align' in parent_css_props and 'text-align' not in css_props:
            css_props['text-align'] = parent_css_props['text-align']
        
        # Skip empty elements (but allow containers like div, ul, ol even if they have no direct text)
        if not element.get_text(strip=True) and tag_name not in ['ul', 'ol', 'div', 'span']:
            return current_page_height
        
        # Handle lists
        if tag_name in ['ul', 'ol']:
            # Calculate total left indent: parent's indent + this list's margin-left
            total_left_indent = Pt(0)
            
            # Get parent's left indent from parent_css_props
            if parent_css_props:
                parent_left = parent_css_props.get('margin-left', parent_css_props.get('padding-left'))
                if parent_left and parent_left != '0':
                    parent_indent = css_margin_to_inches(parent_left)
                    if parent_indent:
                        total_left_indent = parent_indent
            
            # Add this list's own margin-left
            list_left = css_props.get('margin-left', css_props.get('padding-left'))
            if list_left and list_left != '0':
                list_indent = css_margin_to_inches(list_left)
                if list_indent:
                    # Add to parent's indent
                    if hasattr(total_left_indent, 'pt') and hasattr(list_indent, 'pt'):
                        total_left_indent = Pt(total_left_indent.pt + list_indent.pt)
                    else:
                        total_left_indent = list_indent
            
            for li in element.find_all('li', recursive=False):
                if li.get_text(strip=True):
                    p = doc.add_paragraph(style='bulleted')
                    
                    # Get list item's own margins/padding
                    li_classes = li.get('class', [])
                    li_css_props = {}
                    if li_classes:
                        for li_class in li_classes:
                            if li_class in css_styles:
                                li_css_props.update(css_styles[li_class])
                    
                    # Calculate final indent for Word
                    # Word model: left_indent = where text wraps, first_line_indent = bullet offset from text
                    # CSS grid margins on <li> don't translate well to Word, so we skip li margin-left
                    # and only use: parent indent + list indent + li padding
                    final_left_indent = total_left_indent
                    first_line_offset = Pt(0)
                    
                    # Skip list item's margin-left (it's for CSS grid layout, doesn't work in Word's linear model)
                    # Instead, only use padding-left to position text relative to bullet
                    li_padding_left = li_css_props.get('padding-left', '0')
                    if li_padding_left and li_padding_left != '0':
                        li_padding = css_margin_to_inches(li_padding_left)
                        if li_padding and hasattr(final_left_indent, 'pt') and hasattr(li_padding, 'pt'):
                            # Text position = parent + list + padding
                            final_left_indent = Pt(final_left_indent.pt + li_padding.pt)
                            # Bullet stays at parent + list, so offset is -padding
                            first_line_offset = Pt(-li_padding.pt)
                        elif li_padding and hasattr(li_padding, 'pt'):
                            first_line_offset = Pt(-li_padding.pt)
                    
                    # Apply calculated indents
                    if hasattr(final_left_indent, 'pt') and final_left_indent.pt != 0:
                        p.paragraph_format.left_indent = final_left_indent
                    else:
                        p.paragraph_format.left_indent = BULLET_LEFT_INDENT
                    
                    if hasattr(first_line_offset, 'pt') and first_line_offset.pt != 0:
                        p.paragraph_format.first_line_indent = first_line_offset
                    else:
                        p.paragraph_format.first_line_indent = BULLET_FIRST_LINE_INDENT
                    
                    current_page_height = check_and_insert_page_break(p, current_page_height)
                    # Merge list item CSS props with parent list CSS props for font size context
                    li_effective_css_props = css_props.copy()
                    li_effective_css_props.update(li_css_props)
                    process_text_with_hyperlinks(p, li, strip_leading=True, css_styles=css_styles, parent_css_props=li_effective_css_props)
                    # Apply styles from list item classes
                    li_classes = li.get('class', [])
                    if li_classes:
                        effective_class = resolve_descendant_style(li, li_classes[0], css_styles)
                        if effective_class in css_styles:
                            for run in p.runs:
                                apply_style(run, effective_class, css_styles)
                    if li.get('aria-label'):
                        add_hidden_text(p.add_run(), f" | aria-label: {li['aria-label']}")
                    current_page_height = estimate_paragraph_height(p)
            return current_page_height
        
        # Handle divs - check CSS for layout behavior
        if tag_name == 'div':
            # Check if this is a layout container (grid/flex)
            is_layout_container = 'display' in css_props and ('grid' in css_props['display'] or 'flex' in css_props['display'])
            
            if is_layout_container:
                # For grid/flex containers, process children in order
                # Find child elements that contain spans/links (likely the main content)
                # and elements that might be dates
                main_content = None
                date_element = None
                
                # Find the first child div/span that contains spans or links (main content area)
                # Also check if it has text content directly
                for child in element.find_all(['div', 'span'], recursive=False):
                    if child.find(['span', 'a'], recursive=False) or child.get_text(strip=True):
                        # Check if it's not a date element
                        if not is_date_element(child):
                            main_content = child
                            break
                
                # Find element that might contain dates using centralized pattern matching
                for child in element.find_all(['span', 'div'], recursive=False):
                    if is_date_element(child):
                        date_element = child
                        break
                
                # If we found main_content, process it as a structured entry
                if main_content:
                    # This is a structured entry (employment, education, etc.)
                    # Get paragraph style from container classes
                    para_style_name = None
                    for class_name in classes:
                        if class_name in word_styles:
                            word_style = doc.styles[word_styles[class_name]]
                            if hasattr(word_style, 'paragraph_format'):
                                para_style_name = word_styles[class_name]
                                break
                    
                    if para_style_name:
                        p = doc.add_paragraph(style=para_style_name)
                    else:
                        p = doc.add_paragraph()
                    
                    # Clear all inherited tab stops first (before applying spacing or adding content)
                    # This removes default 0.5" tab stops that Word adds automatically
                    # Must clear before adding our custom tab stop
                    # Clear tab stops on the paragraph format (not the style) to override style defaults
                    p.paragraph_format.tab_stops.clear_all()
                    
                    # Apply container spacing
                    if 'margin-top' in css_props:
                        margin_top = css_margin_to_inches(css_props['margin-top'])
                        if margin_top:
                            p.paragraph_format.space_before = margin_top
                    
                    # Set up tab for dates if present
                    date_text = ""
                    if date_element:
                        date_text = date_element.get_text(strip=True)
                        # Right-align dates at the right edge of text area
                        p.paragraph_format.tab_stops.clear_all()
                        right_tab_pos = get_right_tab_stop_position(doc, p)
                        p.paragraph_format.tab_stops.add_tab_stop(right_tab_pos, WD_TAB_ALIGNMENT.RIGHT)
                    
                    # Process main_content children (spans, links) with separators
                    # Process in document order to preserve structure
                    first_child = True
                    for child in main_content.children:
                        # Skip comments and whitespace-only text nodes
                        if isinstance(child, Comment):
                            continue
                        if isinstance(child, str) and not child.strip():
                            continue
                        if not hasattr(child, 'name'):
                            continue
                        # Only process spans and links (skip other elements)
                        if child.name not in ['span', 'a']:
                            continue
                        if not child.get_text(strip=True):
                            continue
                        if not first_child:
                            p.add_run(' | ')
                        first_child = False
                        
                        run_count_before = len(p.runs)
                        # Get child's CSS props for font size context
                        child_css_props = get_effective_css_props(child, css_styles)
                        # Merge with container CSS props
                        child_effective_css_props = css_props.copy()
                        child_effective_css_props.update(child_css_props)
                        process_text_with_hyperlinks(p, child, False, css_styles, parent_css_props=child_effective_css_props)
                        newly_added_runs = p.runs[run_count_before:]
                        
                        # Apply styles from child classes
                        child_classes = child.get('class', [])
                        if child_classes:
                            effective_class = resolve_descendant_style(child, child_classes[0], css_styles)
                            if effective_class in word_styles:
                                try:
                                    char_style = doc.styles[word_styles[effective_class]]
                                    if hasattr(char_style, 'font'):
                                        for run in newly_added_runs:
                                            if char_style.font.size:
                                                run.font.size = char_style.font.size
                                            if hasattr(char_style.font, 'bold'):
                                                run.font.bold = char_style.font.bold
                                            if hasattr(char_style.font, 'italic'):
                                                run.font.italic = char_style.font.italic
                                            if hasattr(char_style.font, 'color') and char_style.font.color.rgb:
                                                run.font.color.rgb = char_style.font.color.rgb
                                except (KeyError, AttributeError):
                                    if effective_class in css_styles:
                                        for run in newly_added_runs:
                                            apply_style(run, effective_class, css_styles)
                            elif effective_class in css_styles:
                                for run in newly_added_runs:
                                    apply_style(run, effective_class, css_styles)
                    
                    # Add date
                    if date_text:
                        p.add_run('\t')
                        date_run = p.add_run(date_text)
                        # Apply styles from date element classes
                        if date_element:
                            date_classes = date_element.get('class', [])
                            if date_classes:
                                effective_class = resolve_descendant_style(date_element, date_classes[0], css_styles)
                                if effective_class in css_styles:
                                    apply_style(date_run, effective_class, css_styles)
                    
                    current_page_height = check_and_insert_page_break(p, current_page_height)
                    current_page_height = estimate_paragraph_height(p)
                    
                    # Process all other direct children (dates, lists, nested divs, etc.)
                    # Skip main_content since we already processed it
                    for child in element.find_all(['div', 'span', 'ul', 'ol'], recursive=False):
                        if child != main_content and child != date_element:  # Don't reprocess what we already handled
                            current_page_height = process_element(child, parent_css_props=css_props)
                    
                    return current_page_height
                else:
                    # Layout container but no main_content found - process all children generically
                    # This handles cases where the structure is different than expected
                    for child in element.find_all(['div', 'span', 'a', 'ul', 'ol'], recursive=False):
                        current_page_height = process_element(child, parent_css_props=css_props)
                    return current_page_height
            else:
                # Non-layout div - check if it has block children
                block_children = element.find_all(['div', 'ul', 'ol'], recursive=False)
                has_block_children = len(block_children) > 0
                
                # Get direct text (not from children)
                direct_text = ''.join(str(node) for node in element.children if isinstance(node, str))
                has_direct_text = direct_text.strip()
                
                if has_block_children:
                    # Has block children - process them directly (they'll create their own paragraphs)
                    # Don't create a paragraph for the parent
                    # Pass parent CSS props so children can inherit text-align and other properties
                    for child in block_children:
                        current_page_height = process_element(child, parent_css_props=css_props)
                    # Also process any inline children (spans, links) if there's no direct text
                    if not has_direct_text:
                        inline_children = element.find_all(['span', 'a'], recursive=False)
                        if inline_children:
                            # Create a paragraph for inline children
                            para_style_name = None
                            for class_name in classes:
                                if class_name in word_styles:
                                    word_style = doc.styles[word_styles[class_name]]
                                    if hasattr(word_style, 'paragraph_format'):
                                        para_style_name = word_styles[class_name]
                                        break
                            
                            if para_style_name:
                                p = doc.add_paragraph(style=para_style_name)
                            else:
                                p = doc.add_paragraph()
                            
                            # Apply text alignment from CSS
                            text_align = css_props.get('text-align', '').lower()
                            if 'center' in text_align:
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            elif 'right' in text_align:
                                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            
                            current_page_height = check_and_insert_page_break(p, current_page_height)
                            for child in inline_children:
                                # Get child's CSS props for font size context
                                child_css_props = get_effective_css_props(child, css_styles)
                                # Merge with parent CSS props
                                child_effective_css_props = css_props.copy()
                                child_effective_css_props.update(child_css_props)
                                process_text_with_hyperlinks(p, child, False, css_styles, parent_css_props=child_effective_css_props)
                                child_classes = child.get('class', [])
                                if child_classes:
                                    effective_class = resolve_descendant_style(child, child_classes[0], css_styles)
                                    if effective_class in css_styles:
                                        newly_added = p.runs[-1:] if p.runs else []
                                        for run in newly_added:
                                            apply_style(run, effective_class, css_styles)
                            current_page_height = estimate_paragraph_height(p)
                    return current_page_height
                else:
                    # No block children - create paragraph and process content directly
                    # Check for text-align in CSS
                    text_align = css_props.get('text-align', '').lower()
                    
                    # Get paragraph style
                    para_style_name = None
                    for class_name in classes:
                        if class_name in word_styles:
                            word_style = doc.styles[word_styles[class_name]]
                            if hasattr(word_style, 'paragraph_format'):
                                para_style_name = word_styles[class_name]
                                break
                    
                    if para_style_name:
                        p = doc.add_paragraph(style=para_style_name)
                    else:
                        p = doc.add_paragraph()
                    
                    # Apply text alignment
                    if 'center' in text_align:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif 'right' in text_align:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    # Apply spacing from CSS
                    if 'margin-top' in css_props or 'padding-top' in css_props:
                        margin_top = css_margin_to_inches(css_props.get('margin-top') or css_props.get('padding-top'))
                        if margin_top:
                            p.paragraph_format.space_before = margin_top
                    else:
                        p.paragraph_format.space_before = Pt(0)
                    
                    if 'margin-bottom' in css_props or 'padding-bottom' in css_props:
                        margin_bottom = css_margin_to_inches(css_props.get('margin-bottom') or css_props.get('padding-bottom'))
                        if margin_bottom:
                            p.paragraph_format.space_after = margin_bottom
                    else:
                        p.paragraph_format.space_after = Pt(0)
                    
                    current_page_height = check_and_insert_page_break(p, current_page_height)
                    process_text_with_hyperlinks(p, element, False, css_styles, parent_css_props=css_props)
                    # Apply styles from element classes to runs
                    for class_name in classes:
                        effective_class = resolve_descendant_style(element, class_name, css_styles)
                        if effective_class in css_styles:
                            for run in p.runs:
                                apply_style(run, effective_class, css_styles)
                    
                    if element.get('aria-label'):
                        add_hidden_text(p.add_run(), f" | aria-label: {element['aria-label']}")
                    
                    current_page_height = estimate_paragraph_height(p)
                    return current_page_height
        
        # Handle spans - these are inline, so they should be added to parent paragraph
        if tag_name == 'span' and parent_paragraph:
            process_text_with_hyperlinks(parent_paragraph, element, False, css_styles, parent_css_props=css_props)
            # Apply styles
            for class_name in classes:
                effective_class = resolve_descendant_style(element, class_name, css_styles)
                if effective_class in css_styles:
                    newly_added_runs = parent_paragraph.runs[-1:] if parent_paragraph.runs else []
                    for run in newly_added_runs:
                        apply_style(run, effective_class, css_styles)
            return current_page_height
        
        return current_page_height
    
    # Process all top-level elements in body
    for element in body.find_all(['div'], recursive=False):
        current_page_height = process_element(element)
    
    # Save
    if output_file is None:
        output_file = Path(html_file).resolve().with_suffix('.docx')
    
    doc.save(str(output_file))
    print(f"Converted to: {output_file}")


if __name__ == '__main__':
    print(scriptname + "\n" + version)
    
    DEFAULT_HTML_FILE = str(SCRIPT_DIR / 'acoven.html')
    
    if len(sys.argv) < 2:
        html_file = DEFAULT_HTML_FILE
        print(f"Using default file: {html_file}")
    else:
        html_file = sys.argv[1]
    
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    convert_html_to_docx(html_file, output_file)

# End, html_to_docx_ats.py